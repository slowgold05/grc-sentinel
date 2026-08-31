from datetime import datetime
from io import BytesIO

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pydantic import BaseModel

from ruleset.generation.models import GeneratedStatement


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)


class EvidenceReference(BaseModel):
    test_id: str
    status: str
    control_ids: list[str]
    tested_at: datetime


def _set_cell_width(cell, width_dxa: int) -> None:
    properties = cell._tc.get_or_add_tcPr()
    width = properties.find(qn("w:tcW"))
    if width is None:
        width = OxmlElement("w:tcW")
    width.set(qn("w:w"), str(width_dxa))
    width.set(qn("w:type"), "dxa")
    if width.getparent() is None:
        properties.append(width)


def _set_table_geometry(table, widths: tuple[int, ...]) -> None:
    properties = table._tbl.tblPr
    for name, value in (("w:tblW", sum(widths)), ("w:tblInd", 120)):
        element = properties.find(qn(name))
        if element is None:
            element = OxmlElement(name)
            properties.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")
    for column, width in zip(table._tbl.tblGrid.gridCol_lst, widths, strict=True):
        column.set(qn("w:w"), str(width))


def export_policy_docx(
    *,
    title: str,
    company_name: str,
    policy_type: str,
    generated_at: datetime,
    ruleset_version: str,
    statements: list[GeneratedStatement],
    evidence: list[EvidenceReference] | None = None,
) -> bytes:
    """Build an auditor-readable policy with a control traceability appendix."""
    document = Document()
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    section.top_margin = section.right_margin = section.bottom_margin = section.left_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name, normal.font.size = "Calibri", Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for style_name, size, color, before, after in (
        ("Title", 23, DARK_BLUE, 0, 8),
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
    ):
        style = document.styles[style_name]
        style.font.name, style.font.size, style.font.color.rgb = "Calibri", Pt(size), color
        style.paragraph_format.space_before, style.paragraph_format.space_after = Pt(before), Pt(after)

    document.add_heading(title, 0)
    subtitle = document.add_paragraph(f"{company_name} | {policy_type}")
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle.runs[0].font.size = Pt(13)
    subtitle.runs[0].font.color.rgb = RGBColor(89, 89, 89)
    for label, value in (
        ("Generated", generated_at.isoformat()),
        ("Ruleset", ruleset_version),
        ("Status", "Draft - professional review required"),
    ):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.add_run(f"{label}: ").bold = True
        paragraph.add_run(value)

    document.add_heading("Policy Statements", level=1)
    for index, statement in enumerate(statements, start=1):
        document.add_heading(f"Statement {index}", level=2)
        document.add_paragraph(statement.text)

    document.add_page_break()
    document.add_heading("Traceability Appendix", level=1)
    document.add_paragraph(
        "Each statement below is linked to controls accepted by deterministic citation verification."
    )
    table = document.add_table(rows=1, cols=3)
    table.autofit = False
    table.style = "Table Grid"
    widths = (1080, 5400, 2880)
    _set_table_geometry(table, widths)
    for cell, width, label in zip(table.rows[0].cells, widths, ("Statement", "Text", "Controls"), strict=True):
        _set_cell_width(cell, width)
        cell.text = label
        cell.paragraphs[0].runs[0].bold = True
    for index, statement in enumerate(statements, start=1):
        cells = table.add_row().cells
        for cell, width in zip(cells, widths, strict=True):
            _set_cell_width(cell, width)
        cells[0].text = str(index)
        cells[1].text = statement.text
        cells[2].text = ", ".join(statement.control_ids)

    if evidence:
        document.add_heading("Control Evidence", level=1)
        evidence_table = document.add_table(rows=1, cols=4)
        evidence_table.style = "Table Grid"
        for cell, label in zip(
            evidence_table.rows[0].cells,
            ("Test", "Verdict", "Controls", "Tested"),
            strict=True,
        ):
            cell.text = label
            cell.paragraphs[0].runs[0].bold = True
        for item in evidence:
            cells = evidence_table.add_row().cells
            for cell, value in zip(
                cells,
                (
                    item.test_id,
                    item.status.upper(),
                    ", ".join(item.control_ids),
                    item.tested_at.isoformat(),
                ),
                strict=True,
            ):
                cell.text = value

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Generated draft | GRC Sentinel").font.size = Pt(9)
    output = BytesIO()
    document.save(output)
    return output.getvalue()
