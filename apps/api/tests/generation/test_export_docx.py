from datetime import UTC, datetime
from io import BytesIO

from docx import Document

from ruleset.generation.export_docx import EvidenceReference, export_policy_docx
from ruleset.generation.models import GeneratedStatement


def test_exports_policy_and_traceability_appendix() -> None:
    content = export_policy_docx(
        title="Access Control Policy",
        company_name="Example Co",
        policy_type="Access Control",
        generated_at=datetime(2026, 8, 30, tzinfo=UTC),
        ruleset_version="hipaa-v2",
        statements=[GeneratedStatement(text="Administrators use MFA.", control_ids=["IA-2"])],
        evidence=[
            EvidenceReference(
                test_id="github-org-mfa-v1",
                status="pass",
                control_ids=["IA-2"],
                tested_at=datetime(2026, 8, 30, tzinfo=UTC),
            )
        ],
    )
    document = Document(BytesIO(content))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "Access Control Policy" in text
    assert "professional review required" in text
    assert document.tables[0].rows[1].cells[2].text == "IA-2"
    assert document.tables[1].rows[1].cells[0].text == "github-org-mfa-v1"
